"""
Build docs/API_REFERENCE.docx from docs/API_REFERENCE.md.

A lightweight Markdown -> Word converter tuned for this document: it handles
headings (#..####), paragraphs, bold (**), inline code (`code`), fenced code
blocks (```), bullet lists (- / >), and GitHub-style pipe tables. Run with:

    python scripts/build_api_docx.py
"""

import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
SRC = os.path.join(ROOT, "docs", "API_REFERENCE.md")
OUT = os.path.join(ROOT, "docs", "API_REFERENCE.docx")

MONO = "Consolas"
BODY = "Calibri"
CODE_BG = "F2F2F2"
CODE_COLOR = RGBColor(0xB0, 0x2A, 0x37)  # inline code red
LINK_HEADING = RGBColor(0x1F, 0x4E, 0x79)


def shade(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def add_inline(paragraph, text):
    """Render **bold** and `code` inside a paragraph."""
    # Split keeping delimiters for **bold** and `code`.
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO
            run.font.size = Pt(9.5)
            run.font.color.rgb = CODE_COLOR
        else:
            paragraph.add_run(part)


def add_code_block(doc, lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(6)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)
    ppr = para._p.get_or_add_pPr()
    shd = ppr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODE_BG)
    ppr.append(shd)
    text = "\n".join(lines)
    run = para.add_run(text)
    run.font.name = MONO
    run.font.size = Pt(9)


def add_table(doc, rows):
    header = rows[0]
    body = rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, text in enumerate(header):
        hdr[i].text = ""
        add_inline(hdr[i].paragraphs[0], text)
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
    for row in body:
        cells = table.add_row().cells
        for i in range(len(header)):
            cells[i].text = ""
            val = row[i] if i < len(row) else ""
            add_inline(cells[i].paragraphs[0], val)
    doc.add_paragraph()


def parse_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_separator(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def main():
    with open(SRC, encoding="utf-8") as fh:
        raw = fh.read().split("\n")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(10.5)

    i = 0
    n = len(raw)
    while i < n:
        line = raw[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            block = []
            while i < n and not raw[i].strip().startswith("```"):
                block.append(raw[i])
                i += 1
            add_code_block(doc, block)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < n and is_separator(raw[i + 1]):
            rows = [parse_table_row(stripped)]
            i += 2  # skip header + separator
            while i < n and raw[i].strip().startswith("|"):
                rows.append(parse_table_row(raw[i]))
                i += 1
            add_table(doc, rows)
            continue

        # Horizontal rule
        if stripped == "---":
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = re.sub(r"[#`*]", "", m.group(2)).strip()
            if level == 1:
                h = doc.add_heading(level=0)
                add_inline(h, text)
            else:
                h = doc.add_heading(level=min(level - 1, 4))
                h.runs.clear() if h.runs else None
                # add_heading created an empty run set; set text via inline
                for r in list(h.runs):
                    r.text = ""
                add_inline(h, text)
                for r in h.runs:
                    r.font.color.rgb = LINK_HEADING
            i += 1
            continue

        # Bullets / blockquote
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            i += 1
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            add_inline(p, stripped[2:])
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # Numbered TOC / list
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(2))
            i += 1
            continue

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
